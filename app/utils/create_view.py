import os
import re
from abc import ABC, abstractmethod
from glob import glob
from tkinter import font
from typing import List

import pytz
from bs4 import BeautifulSoup
from docx import Document
from justifytext import justify
from PIL import Image, ImageDraw, ImageFont
from requests import get
from requests.exceptions import ConnectionError

from app.__colors__ import blue, grey, light_blue, white
from app.frames.clima import Station
from app.frames.messagebox import box
from app.utils import Fonts, logger
from app.utils.date_utils import TODAY, TOMORROW, YESTERDAY, date2str, tomorrow2str
from app.utils.taf_model import TAF
from app.utils.winds_model import Wind

fonts = Fonts()


def view_creator(func):
    template_path = "assets/img/template.png"

    dirname = "images/output"
    if not os.path.exists(dirname):
        logger.info(f"Path {dirname} doesn't exists. Creating it.")
        os.makedirs(dirname)

    def wrapper(*args, **kwargs):
        logger.info(f"Opening template image to create {args[0]}.")
        img = Image.open(template_path)
        draw = ImageDraw.Draw(img)
        # img, draw = func(img=img, draw=draw, **kwargs)
        result = func(img=img, draw=draw, **kwargs)
        if result == "ok":
            logger.info(f"Saving {args[0]}.")
            img = img.resize((1200, 900))
            img.save("images/output/" + args[0])
            return False
        elif result == "no":
            return False
        return True

    return wrapper


def _make_title(draw: ImageDraw.Draw, text: str, x=0, y=90, color=light_blue):
    logger.info(f"Making image title {text[:10]}...")
    text = text.center(46, " ")
    draw.text((x, y), text, font=fonts.title, fill=color)


def _make_subtitle(draw: ImageDraw.Draw, text: str, x=400, y=210):
    logger.info(f"Making image subtitle {text[:10]}...")
    text = text.center(40, " ")
    draw.text((x, y), text, font=fonts.subtitle, fill=light_blue)


def _make_text(
    draw: ImageDraw.Draw,
    text: str,
    x=200,
    y=325,
    color=blue,
    font=fonts.text,
    just=True,
):
    logger.info(f"Making image text {text[:10]}...")
    if just is True:
        ltext = justify(text, 70)
        draw.text((x, y), "\n".join(ltext), font=font, fill=color)
    else:
        ltext = text.split("\n")
        draw.text((x, y), text, font=font, fill=color)

    return len(ltext * 50)


###########################################################################
############################## CREATE MAP VIEW ############################
###########################################################################


@view_creator
def create_map_img(*args, **kwargs):
    img = kwargs.get("img")
    map_img_path = kwargs.get("map").name
    draw = kwargs.get("draw")

    _make_title(draw, "Meteorología Aeronáutica")
    _make_subtitle(draw, date2str().capitalize())
    if TODAY.hour >= 11:
        logger.info(f"Adding 'ACTUALIZADO' label to report.")
        _make_text(draw, "ACTUALIZADO", x=210, y=270, color=blue)

    logger.info(f"Adding SIGWX Map to image.")
    map_img = Image.open(map_img_path)
    map_img = map_img.resize((2000, 1294))
    img.paste(map_img, (200, 335))

    return "ok"


###########################################################################
############################# CREATE TREND VIEW ###########################
###########################################################################


class TrendText:
    def __init__(self, path):
        self.document = Document(path)
        self._text_from_docx()

    def _text_from_docx(self):
        self.paragraphs = [
            p.text for p in self.document.paragraphs if not re.match(r"^\s*$", p.text)
        ]
        self.title = self.paragraphs[0]
        self.subtitle = self.paragraphs[1]
        self.valid = self.paragraphs[2]
        self.general = [p for p in self.paragraphs[3:5]]
        self.aerodromes = [p for p in self.paragraphs[5:-1]]


@view_creator
def create_trend01(*args, **kwargs):
    draw = kwargs.get("draw")
    docx = kwargs.get("docx")

    if docx is None:
        logger.info(f"create_trend01(): docx file is NONE, returning.")
        return
    logger.info(f"create_trend01(): Obtaining text from docx file.")
    text = TrendText(docx)
    _make_title(draw, text.title)
    _make_subtitle(draw, text.subtitle)
    _ = _make_text(draw, text.valid.center(70))

    y_text = 500
    _ = _make_text(draw, text.general[0], y=y_text, color=light_blue)
    y_text += 75
    pxls = _make_text(draw, text.general[1], y=y_text)
    y_text += pxls + 75
    _ = _make_text(draw, text.aerodromes[0], y=y_text, color=light_blue)
    y_text += 75
    _ = _make_text(draw, text.aerodromes[1], y=y_text)

    return "ok"


@view_creator
def create_trend02(*args, **kwargs):
    draw = kwargs.get("draw")
    docx = kwargs.get("docx")

    if docx is None:
        logger.info(f"create_trend02(): docx file is NONE, returning.")
        return
    logger.info(f"create_trend02(): Obtaining text from docx file.")
    text = TrendText(kwargs.get("docx"))
    _make_title(draw, text.title)
    _make_subtitle(draw, text.subtitle)
    _ = _make_text(draw, text.valid.center(70))

    y_text = 450
    _ = _make_text(draw, text.aerodromes[2], y=y_text, color=light_blue)
    y_text += 75
    pxls = _make_text(draw, text.aerodromes[3], y=y_text)
    y_text += pxls + 65
    _ = _make_text(draw, text.aerodromes[4], y=y_text, color=light_blue)
    y_text += 75
    pxls = _make_text(draw, text.aerodromes[5], y=y_text)
    y_text += pxls + 65
    _make_text(draw, text.aerodromes[6], y=y_text, color=light_blue)
    y_text += 75
    _make_text(draw, text.aerodromes[7], y=y_text)

    return "ok"


###########################################################################
############################# CREATE VASH VIEW ############################
###########################################################################

vash_text = "Modelos de dispersión de ceniza volcánica, validez hasta las 12:00Z del {}. Indica la dispersión de la ceniza volcánica para los 3350, 4000 y 5000 msnm, en erupciones superiores a los 500 m."


def _paste_vash_img(
    img: Image,
    img_num: int,
    name: str,
    dirname: str,
    img_size=(850, 1055),
    paste_pos=(350, 550),
):
    found = False
    for fmt in [".png", ".jpg", ".gif", ".jpeg", ".bmp"]:
        try:
            img_path = f"images/volcanoes/{dirname}/image{img_num}{fmt}"
            logger.info(f"Try to open volcanic ash image: {img_path}")
            print(img_path)
            ash_img = Image.open(img_path)
        except FileNotFoundError as e:
            logger.debug(f"{e}")
            continue
        else:
            logger.info(f"Volcanic ash image found: {img_path}")
            ash_img = ash_img.resize(img_size)
            img.paste(ash_img, paste_pos)
            found = True
            break

    if not found:
        msg = f"No se encuentra la imagen {img_num} del Volcán {name.title()}."
        logger.debug(f"Raising FileNotFoundError: {msg}")
        raise FileNotFoundError(msg)


@view_creator
def create_volcanic_ash(*args, **kwargs):
    img = kwargs.get("img")
    draw = kwargs.get("draw")
    name = kwargs.get("name")
    dirname = kwargs.get("dir")

    _make_title(draw, "Dispersión de Ceniza")
    _make_subtitle(draw, f"Volcán {name}")
    _make_text(draw, vash_text.format(tomorrow2str()))

    box_params = ["okcancel", "Faltan imágenes."]
    with_errors = False
    images = glob(f"images/volcanoes/{dirname}/*")

    if len(images) == 6:
        logger.warning(
            f"less images than expected found in images/volcanoes/{dirname}. Only one will be pasted."
        )
        try:
            _paste_vash_img(img, 1, name, dirname, paste_pos=(775, 550))
        except FileNotFoundError as e:
            result = box(*box_params, e)
            if result:
                logger.info("User choose to continue creating view with errors.")
                with_errors = True
            else:
                logger.info("User choose to stop creating report.")
                return
    elif len(images) in [4, 7]:
        logger.warning(
            f"correct number of images found in images/volcanoes/{dirname}. Two will be pasted."
        )
        try:
            _paste_vash_img(img, 1, name, dirname)
        except FileNotFoundError as e:
            result = box(*box_params, e)
            if result:
                logger.info("User choose to continue creating view with errors.")
                with_errors = True
            else:
                logger.info("User choose to stop creating report.")
                return

        try:
            _paste_vash_img(
                img, 2, name, dirname, img_size=(850, 797), paste_pos=(1230, 700)
            )
        except FileNotFoundError as e:
            result = box(*box_params, e)
            if result:
                logger.info("User choose to continue creating view with errors.")
                with_errors = True
            else:
                logger.info("User choose to stop creating report.")
                return
    else:
        result = box(
            *box_params,
            f"No se encuentra la cantidad correcta de imágenes para el volcán {name}. No se creará la imagen.",
        )

    if with_errors:
        result = box(
            "okcancel",
            "Imagen con errores.",
            "¿Desea guardar la imagen para este volcán?",
        )
        if not result:
            logger.info("User choose do not save volcanic ash view.")
            return "no"
    logger.info("User choose save volcanic ash view.")
    return "ok"


###########################################################################
############################## CREATE TAF VIEW ############################
###########################################################################

# BASE_URL_ADDS = "http://www.aviationweather.gov/adds/dataserver_current/httpparam?dataSource=tafs&requestType=retrieve&format=csv&stationString={}&hoursBeforeNow=0"
# BASE_URL_NOAA = "https://tgftp.nws.noaa.gov/data/forecasts/taf/stations/{}.TXT"

BASE_URL_OGIMET = "http://ogimet.com/display_metars2.php?lugar=MROC+MRLB+MRLM+MRPV&tipo=FT&ord=REV&nil=NO&fmt=txt&ano={}&mes={:02d}&day={:02d}&hora={:02d}&anof={}&mesf={:02d}&dayf={:02d}&horaf={:02d}&minf=59&enviar=Ver"


def _process_ogimet_taf(taf: List[str]):
    first_line = taf[0] + " " + taf[1]

    return " ".join([first_line] + taf[2:] if len(taf) > 2 else [first_line])


def _taf_from_ogimet():
    UTC = pytz.utc
    tafs = []
    yestarday = YESTERDAY.astimezone(tz=UTC)
    today = TODAY.astimezone(tz=UTC)
    url = BASE_URL_OGIMET.format(
        yestarday.year,
        yestarday.month,
        yestarday.day,
        yestarday.hour,
        today.year,
        today.month,
        today.day,
        today.hour,
    )

    res = get(url)
    if res.status_code != 200:
        raise ConnectionError

    taf_found = False
    taf = []
    for line in res.text.split("\n"):
        if taf_found and not line.startswith("#"):
            taf.append(line.strip())
            if line.endswith("="):
                taf_found = False
                tafs.append(_process_ogimet_taf(taf))
                taf = []
        if re.match(r"#\s+TAF\s+LARGOS", line):
            taf_found = True

    return tafs


BASE_URL_ADDS = "https://www.aviationweather.gov/taf/data?ids=MROC+MRLB+MRLM+MRPV&format=raw&date=&submit=Get+TAF+data"


def _taf_from_adds():
    tafs = []
    res = get(BASE_URL_ADDS)
    if res.status_code == 404:
        raise ConnectionError
    soup = BeautifulSoup(res.text, "html.parser")

    for taf in soup.find_all("code"):
        tafs.append(taf.getText())

    return tafs


@view_creator
def create_taf(*args, **kwargs):
    draw = kwargs.get("draw")

    _make_title(draw, "Terminal Aerodrome Forecast (TAF)")

    subtitle = "TAF válidos hasta las {}:00Z del {}"
    if TODAY.hour >= 17:
        subtitle = subtitle.format("00", tomorrow2str(days=2))
    elif TODAY.hour >= 11:
        subtitle = subtitle.format("18", tomorrow2str())
    elif TODAY.hour >= 5:
        subtitle = subtitle.format("12", tomorrow2str())
    else:
        subtitle = subtitle.format("06", tomorrow2str())
    _make_text(draw, subtitle, x=330, y=280, color=light_blue, just=False)

    with_errors = False
    try:
        logger.info("Try get TAF from ADDS.")
        tafs = _taf_from_adds()
    except ConnectionError as e:
        logger.error("ConnectionError: {e}.")
        try:
            logger.info("Try get TAF from Ogimet.")
            tafs = _taf_from_ogimet()
        except ConnectionError as e:
            logger.error("ConnectionError: {e}.")
            result = box(
                "okcancel",
                "Error de conexión.",
                "No se puede acceder a los TAF. ¿Desea continuar? Se omitirá esta imagen.",
            )
            if result:
                with_errors = True
            else:
                return

    if with_errors:
        logger.info("User choose do not save TAF view.")
        return "no"
    else:
        logger.info("User choose save TAF view.")
        y_text = 420
        for taf in tafs:
            taf = TAF(taf)
            pxls = _make_text(draw, taf.formated, x=100, y=y_text, just=False)
            y_text += pxls + 35
        return "ok"


###########################################################################
############################# CREATE WINDS VIEW ###########################
###########################################################################


def _draw_winds_table(draw: ImageDraw.Draw):
    # light blue rectangle
    draw.rectangle((400, 450, 2208, 700), fill=light_blue)

    # grey rectangles
    x_left = 400
    for i in range(2):
        x_right = x_left + 452
        draw.rectangle((x_left, 700, x_right, 1606), fill=grey)
        x_left += 904

    # longest vertical lines
    x_left = 400
    for i in range(5):
        draw.line((x_left, 448, x_left, 1608), fill=blue, width=5)
        x_left += 452

    # middle vertical lines
    x_left = 512
    for i in range(4):
        for j in range(3):
            draw.line((x_left, 550, x_left, 1606), fill=blue, width=5)
            x_left += 113
        x_left += 113

    # light blue rectanle horizontal lines
    y_top = 450
    for i in range(3):
        if i == 1:
            y_top += 25
        draw.line((400, y_top, 2208, y_top), fill=blue, width=5)
        y_top += 75

    # longest horizontal lines
    for i in range(7):
        draw.line((148, y_top, 2210, y_top), fill=blue, width=5)
        y_top += 151
    # most left vertical line
    draw.line((150, 700, 150, 1606), fill=blue, width=5)


def _write_winds_table_text(draw: ImageDraw.Draw):
    draw.text((210, 570), "Fecha", font=fonts.table, fill=blue)
    draw.text((180, 640), "Hora UTC", font=fonts.table, fill=blue)
    draw.text((90, 950), "N\n\nI\n\nV\n\nE\n\nL", font=fonts.table, fill=blue)

    note = (
        "Nota: los datos aquí presentados corresponden al promedio de las 6 horas anteriores."
        " La dirección corresponde a la procedencia del viento."
        " CLM indica viento calmo."
    )
    note = justify(note, 84)
    draw.text(
        (150, 1640),
        "\n".join(note),
        font=fonts.note,
        fill=blue,
    )

    levels = [
        " 300 hPa \n33.000 ft",
        " 400 hPa \n25.000 ft",
        " 500 hPa \n20.000 ft",
        " 700 hPa \n10.000 ft",
        " 850 hPa \n 5.000 ft",
        " 925 hPa \n 3.000 ft",
    ]
    y_top = 730
    for level in levels:
        draw.text((170, y_top), level, font=fonts.table, fill=blue)
        y_top += 151


BASE_WINDS_U_URL = (
    "http://wrf1-5.imn.ac.cr/modelo/informe_aeronautico/salidas/informe_{}_U.txt"
)
BASE_WINDS_V_URL = (
    "http://wrf1-5.imn.ac.cr/modelo/informe_aeronautico/salidas/informe_{}_V.txt"
)


def _sanitize_str(s):
    s = re.sub(r"\s{2,}", " ", s)
    s = s.strip()

    return s


def _process_response(u_text, v_text):
    indexes = (6, 9)
    u_temp = u_text.text.split("\n")
    u_temp = [u_temp[indexes[0]]] + u_temp[indexes[1] : -1]
    v_temp = v_text.text.split("\n")
    v_temp = [v_temp[indexes[0]]] + v_temp[indexes[1] : -1]

    u, v = [], []
    for u_el, v_el in zip(u_temp, v_temp):
        u_el = _sanitize_str(u_el)
        v_el = _sanitize_str(v_el)

        u.append(u_el.upper())
        v.append(v_el.upper())

    return u, v


def _get_winds_data():
    winds = {}
    for stn in ["MROC", "MRLB", "MRLM", "MRPV"]:
        u_url = BASE_WINDS_U_URL.format(stn)
        v_url = BASE_WINDS_V_URL.format(stn)
        try:
            logger.info(f"Try get winds data for station {stn}.")
            u_res = get(u_url)
            v_res = get(v_url)
        except ConnectionError as e:
            logger.error(f"ConnectionError with winds source for {stn}: {e}.")
            result = box(
                "okcancel",
                f"Error de conexión.",
                f"No se puede conectar la petición de los datos de vientos en altura de {stn}. ¿Desea continuar?",
            )
            if not result:
                raise
            continue
        else:
            u, v = _process_response(u_res, v_res)
            winds[stn] = Wind(u, v)

    logger.info("All winds obtained succesfully.")
    return winds


def _write_winds_on_table(draw: ImageDraw.Draw, winds: dict):
    x = 400
    y = 450
    for stn, wind in winds.items():
        date = TODAY
        _make_text(draw, stn, x=x + 130, y=y, color=white, font=fonts.title)
        _x = 25
        for time in wind.hours:
            _y = 120
            _make_text(
                draw,
                "{:02d}".format(date.day),
                x=x + _x + 10,
                y=y + _y,
                color=white,
                font=fonts.table,
            )
            _y += 75
            _make_text(draw, time, x=x + _x, y=y + _y, color=white, font=fonts.table)
            _y += 85
            for level in [300, 400, 500, 700, 850, 925]:
                text = wind.values(time, level)
                _make_text(
                    draw,
                    text,
                    x=x + _x,
                    y=y + _y,
                    color=blue,
                    just=False,
                    font=fonts.table,
                )
                _y += 151
            _x += 113
            date = TOMORROW
        x += 450


@view_creator
def create_winds(*args, **kwargs):
    draw = kwargs.get("draw")

    # get the winds from internet
    try:
        winds = _get_winds_data()
    except ConnectionError:
        return

    _make_title(draw, "Vientos en Altura")
    _make_subtitle(draw, "Dirección y Velocidad (kt)")
    _make_text(
        draw,
        f"Válido hasta las {'12'}:00Z del {tomorrow2str()}",
        x=400,
        y=345,
    )

    _draw_winds_table(draw)
    _write_winds_table_text(draw)
    _write_winds_on_table(draw, winds)

    logger.info("Saving winds table view.")
    return "ok"


###########################################################################
############################# CREATE CLIMA VIEW ###########################
###########################################################################


def _draw_clima_table(draw: ImageDraw.Draw):
    draw.rectangle((200, 350, 2200, 550), fill=light_blue)

    # draw vertical lines
    x_left = 200
    for i in range(5):
        if i == 1:
            x_left += 200
        draw.line((x_left, 350, x_left, 950), fill=blue, width=5)
        x_left += 450

    # draw horizontal lines
    x_left = 198
    y_top = 350
    for i in range(6):
        if i == 1:
            y_top += 100
        draw.line((x_left, y_top, 2202, y_top), fill=blue, width=5)
        y_top += 100


def _write_clima_table_text(
    draw: ImageDraw.Draw,
    clima: List[Station],
):
    titles = [
        "  Estación   \nMeteorológica",
        "  T. Máxima  \n     (°)     ",
        "  T. Mínima  \n     (°)     ",
        "Precipitación\n    (mm)     ",
    ]

    x_left = 340
    for title in titles:
        if titles.index(title) == 1:
            x_left += 90
        if titles.index(title) == 3:
            x_left += 10
        _make_text(draw, title, color=white, x=x_left, y=390, just=False)
        x_left += 450

    y_top = 575
    stations = [
        "Juan Santamaría (MROC)",
        "Daniel Oduber (MRLB)",
        "Limón (MRLM)",
        "Tobías Bolaños (MRPV)",
    ]
    for stn in stations:
        stn = stn.center(25, " ")
        _make_text(draw, stn, color=blue, x=220, y=y_top, just=False, font=fonts.table)
        y_top += 100

    x_left = 1010
    y_top = 575
    for stn in clima:
        tmax, tmin, prec = stn.get_values()
        _make_text(
            draw,
            tmax.center(5, " "),
            x=x_left,
            y=y_top,
            just=False,
            color=blue,
            font=fonts.table,
        )
        _make_text(
            draw,
            tmin.center(5, " "),
            x=x_left + 450,
            y=y_top,
            just=False,
            color=blue,
            font=fonts.table,
        )
        _make_text(
            draw,
            prec.center(5, " "),
            x=x_left + 900,
            y=y_top,
            just=False,
            color=blue,
            font=fonts.table,
        )
        y_top += 100


def _write_ephemeris(
    img: Image,
    draw: ImageDraw.Draw,
    data=("00:00 AM", "00:00 AM"),
):
    _make_text(
        draw,
        "Salida y Puesta del Sol",
        color=blue,
        just=False,
        x=250,
        y=1000,
        font=fonts.subtitle,
    )
    _make_text(draw, "Salida de mañana", color=blue, just=False, x=250, y=1100)
    _make_text(draw, "Puesta de hoy", color=blue, just=False, x=775, y=1100)

    sunrise = Image.open("assets/img/sunrise.png")
    sunset = Image.open("assets/img/sunset.png")
    img.paste(sunrise, (250, 1180))
    img.paste(sunset, (735, 1180))

    draw.rectangle((250, 1500, 1200, 1600), fill=light_blue)
    x_left = 370
    for d in data:
        _make_text(draw, d, color=white, x=x_left, y=1525)
        x_left += 500


def _write_user_data(draw: ImageDraw.Draw, data=("Name", "email@email.com")):
    text = "{}\n{}\n{}\n{}@imn.ac.cr\n{}".format(
        data[0],
        "Meteorología Aeronáutica (IMN)",
        "Aeropuerto Int. Tobías Bolaños",
        data[1],
        "Telefax: (+506) 2232-2071\nWeb (IMN): www.imn.ac.cr",
    )

    _make_text(draw, text.strip(), color=blue, x=1280, y=1200, just=False)


@view_creator
def create_clima(*args, **kwargs):
    img = kwargs.get("img")
    draw = kwargs.get("draw")
    clima = kwargs.get("clima")
    ephemeris = kwargs.get("ephemeris")
    user = kwargs.get("user")

    yesterday = date2str(date=YESTERDAY)

    _make_title(draw, "Datos Climatológicos")
    _make_subtitle(draw, yesterday.capitalize())

    _draw_clima_table(draw)
    _write_clima_table_text(draw, clima)
    _write_ephemeris(img, draw, data=ephemeris)
    _write_user_data(draw, data=user)

    logger.info("Saving climatology view.")
    return "ok"
